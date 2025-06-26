#!/usr/bin/env python3
"""
Script to split facility data files into individual JSON files for each record.

This script:
1. Reads all JSON files from the facility_data folder
2. For each input file, extracts individual records
3. Creates a subfolder for each input file
4. Saves each record as a separate JSON file in the corresponding subfolder
"""

import json
import os
import sys
from pathlib import Path
from typing import List, Dict, Any, Union
import argparse

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX output will be skipped.")
    print("Install with: pip install python-docx")

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: pandas not available. Excel output will be skipped.")
    print("Install with: pip install pandas openpyxl")


def load_records_from_file(file_path: Path) -> List[Dict[str, Any]]:
    """
    Load records from a JSON file, handling different structures.

    Args:
        file_path: Path to the JSON file

    Returns:
        List of record dictionaries
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Handle different JSON structures
        if isinstance(data, list):
            # Direct array of records
            return data
        elif isinstance(data, dict):
            if 'records' in data:
                # Records nested under 'records' key
                return data['records']
            elif len(data) == 1 and isinstance(list(data.values())[0], list):
                # Single key with array value
                return list(data.values())[0]
            else:
                # Single record as object
                return [data]
        else:
            print(f"Warning: Unexpected data structure in {file_path}")
            return []

    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in {file_path}: {e}")
        return []
    except Exception as e:
        print(f"Error: Could not read {file_path}: {e}")
        return []


def create_safe_filename(record: Dict[str, Any], index: int, source_filename: str = "") -> str:
    """
    Create a safe filename for the record JSON file using action request number + facility name.

    Args:
        record: The record dictionary
        index: Index of the record in the original file
        source_filename: Name of the source file (used as fallback for facility name)

    Returns:
        Safe filename string in format: ActionRequestNumber_FacilityName.json
    """
    # Get action request number
    action_request = record.get("Action Request Number:", "")
    if not action_request:
        action_request = record.get("Action Request Number", "")

    # Get facility name
    facility_name = record.get("_facility_name", "")
    if not facility_name and source_filename:
        # Use source filename without extension as facility name
        facility_name = Path(source_filename).stem

    # Clean values to make them filesystem-safe
    if action_request:
        safe_action = "".join(c for c in str(action_request).strip() if c.isalnum() or c in ('-', '_')).strip()
    else:
        safe_action = f"Record{index:04d}"

    if facility_name:
        safe_facility = "".join(c for c in str(facility_name).strip() if c.isalnum() or c in ('-', '_')).strip()
    else:
        safe_facility = "UnknownFacility"

    # Create filename: ActionRequestNumber_FacilityName.json
    filename = f"{safe_action}_{safe_facility}.json"

    # Ensure filename is not too long (max 255 characters for most filesystems)
    if len(filename) > 200:
        safe_action = safe_action[:100]
        safe_facility = safe_facility[:90]
        filename = f"{safe_action}_{safe_facility}.json"

    return filename


def split_facility_file(input_file: Path, output_base_dir: Path, create_docx: bool = False, create_excel: bool = False) -> int:
    """
    Split a single facility file into individual record files.

    Args:
        input_file: Path to the input JSON file
        output_base_dir: Base directory for output
        create_docx: Whether to create DOCX files alongside JSON files
        create_excel: Whether to create Excel files alongside JSON files

    Returns:
        Number of records processed
    """
    print(f"Processing {input_file.name}...")

    # Load records from the file
    records = load_records_from_file(input_file)

    if not records:
        print(f"  No records found in {input_file.name}")
        return 0

    # Create output subfolders for different file types
    subfolder_name = input_file.stem  # Filename without extension

    # Create separate folders for JSON, DOCX, and Excel outputs
    json_output_dir = output_base_dir / "json" / subfolder_name
    docx_output_dir = output_base_dir / "docx" / subfolder_name
    excel_output_dir = output_base_dir / "excel" / subfolder_name

    json_output_dir.mkdir(parents=True, exist_ok=True)
    if create_docx:
        docx_output_dir.mkdir(parents=True, exist_ok=True)
    if create_excel:
        excel_output_dir.mkdir(parents=True, exist_ok=True)

    # Process each record
    processed_count = 0
    for index, record in enumerate(records):
        try:
            # Create filename using action request number + facility name
            filename = create_safe_filename(record, index, input_file.name)
            base_filename = filename.rsplit('.', 1)[0]  # Remove .json extension

            json_output_file = json_output_dir / filename
            docx_output_file = docx_output_dir / f"{base_filename}.docx"
            excel_output_file = excel_output_dir / f"{base_filename}.xlsx"

            # Add metadata to the record
            record_with_meta = record.copy()
            record_with_meta['_split_metadata'] = {
                'source_file': input_file.name,
                'record_index': index,
                'split_timestamp': None  # Will be filled at runtime if needed
            }

            # Write the individual JSON record file
            with open(json_output_file, 'w', encoding='utf-8') as f:
                json.dump(record_with_meta, f, indent=2, ensure_ascii=False)

            # Create DOCX file if requested
            docx_success = False
            if create_docx:
                docx_success = create_docx_from_record(record_with_meta, docx_output_file)

            # Create Excel file if requested
            excel_success = False
            if create_excel:
                excel_success = create_excel_from_record(record_with_meta, excel_output_file)

            processed_count += 1

            # Log what was created
            created_files = [f"JSON: {filename}"]
            if create_docx and docx_success:
                created_files.append(f"DOCX: {base_filename}.docx")
            elif create_docx and not docx_success:
                created_files.append("DOCX: Failed")

            if create_excel and excel_success:
                created_files.append(f"Excel: {base_filename}.xlsx")
            elif create_excel and not excel_success:
                created_files.append("Excel: Failed")

            if index < 5:  # Only show first few for brevity
                print(f"    Created: {', '.join(created_files)}")

        except Exception as e:
            print(f"  Error processing record {index}: {e}")
            continue

    format_info_parts = ["JSON"]
    if create_docx and DOCX_AVAILABLE:
        format_info_parts.append("DOCX")
    if create_excel and EXCEL_AVAILABLE:
        format_info_parts.append("Excel")

    format_info = f" ({' + '.join(format_info_parts)} files)"
    print(f"  Created {processed_count} individual files{format_info}")
    print(f"    JSON files: {json_output_dir}")
    if create_docx and DOCX_AVAILABLE:
        print(f"    DOCX files: {docx_output_dir}")
    if create_excel and EXCEL_AVAILABLE:
        print(f"    Excel files: {excel_output_dir}")
    return processed_count


def create_docx_from_record(record: Dict[str, Any], output_file: Path) -> bool:
    """
    Create a DOCX document from a record dictionary.

    Args:
        record: The record dictionary
        output_file: Path where the DOCX file should be saved

    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        # Create a new document
        doc = Document()

        # Add title
        title = doc.add_heading('Mining Maintenance Record', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add action request number as subtitle if available
        action_request = record.get("Action Request Number:", "") or record.get("Action Request Number", "")
        if action_request:
            subtitle = doc.add_heading(f'Action Request: {action_request}', level=1)
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add facility name if available
        facility_name = record.get("_facility_name", "")
        if facility_name:
            facility_para = doc.add_paragraph()
            facility_run = facility_para.add_run(f'Facility: {facility_name}')
            facility_run.bold = True
            facility_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a line break
        doc.add_paragraph()

        # Add all record fields as a table or formatted text
        # Skip metadata fields
        skip_fields = {'_split_metadata', '_facility_name'}

        # Get all non-empty fields
        content_fields = []
        for field_name, field_value in record.items():
            if field_name in skip_fields:
                continue
            if field_value is not None and str(field_value).strip():
                content_fields.append((field_name, field_value))

        if content_fields:
            # Create a table for better formatting
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Set header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'

            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for field_name, field_value in content_fields:
                row_cells = table.add_row().cells

                # Clean field name for display
                display_name = field_name.replace('_', ' ').title()
                if display_name.endswith(':'):
                    display_name = display_name[:-1]

                row_cells[0].text = display_name

                # Handle long text values
                value_str = str(field_value)
                if len(value_str) > 100:
                    # For long text, add line breaks for better readability
                    row_cells[1].text = value_str
                else:
                    row_cells[1].text = value_str

        # Add metadata section if present
        if '_split_metadata' in record:
            doc.add_page_break()
            doc.add_heading('Document Information', level=1)

            metadata = record['_split_metadata']
            meta_para = doc.add_paragraph()
            meta_para.add_run('Source File: ').bold = True
            meta_para.add_run(str(metadata.get('source_file', 'Unknown')))
            meta_para.add_run('\nRecord Index: ').bold = True
            meta_para.add_run(str(metadata.get('record_index', 'Unknown')))

        # Save the document
        doc.save(output_file)
        return True

    except Exception as e:
        print(f"Error creating DOCX file {output_file}: {e}")
        return False


def create_excel_from_record(record: Dict[str, Any], output_file: Path) -> bool:
    """
    Create an Excel file from a record dictionary.

    Args:
        record: The record dictionary
        output_file: Path where the Excel file should be saved

    Returns:
        True if successful, False otherwise
    """
    if not EXCEL_AVAILABLE:
        return False

    try:
        # Create a new workbook with multiple sheets
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:

            # Main record data sheet
            # Skip metadata fields for main data
            skip_fields = {'_split_metadata'}
            main_data = []

            for field_name, field_value in record.items():
                if field_name in skip_fields:
                    continue
                if field_value is not None and str(field_value).strip():
                    # Clean field name for display
                    display_name = field_name.replace('_', ' ').title()
                    if display_name.endswith(':'):
                        display_name = display_name[:-1]

                    main_data.append({
                        'Field': display_name,
                        'Value': str(field_value)
                    })

            if main_data:
                df_main = pd.DataFrame(main_data)
                df_main.to_excel(writer, sheet_name='Record_Data', index=False)

                # Auto-adjust column widths
                worksheet = writer.sheets['Record_Data']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
                    worksheet.column_dimensions[column_letter].width = adjusted_width

            # Metadata sheet if present
            if '_split_metadata' in record:
                metadata = record['_split_metadata']
                meta_data = []
                for key, value in metadata.items():
                    if value is not None:
                        meta_data.append({
                            'Property': key.replace('_', ' ').title(),
                            'Value': str(value)
                        })

                if meta_data:
                    df_meta = pd.DataFrame(meta_data)
                    df_meta.to_excel(writer, sheet_name='Document_Info', index=False)

                    # Auto-adjust column widths for metadata sheet
                    worksheet_meta = writer.sheets['Document_Info']
                    for column in worksheet_meta.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 30)
                        worksheet_meta.column_dimensions[column_letter].width = adjusted_width

        return True

    except Exception as e:
        print(f"Error creating Excel file {output_file}: {e}")
        return False


def main():
    """Main function to process all facility data files."""
    parser = argparse.ArgumentParser(description="Split facility data files into individual JSON, DOCX, and Excel records")
    parser.add_argument("--input-dir",
                       default="data/facility_data",
                       help="Input directory containing facility data files (default: data/facility_data)")
    parser.add_argument("--output-dir",
                       default="data/facility_data/split_records",
                       help="Output directory for split records (default: data/facility_data/split_records)")
    parser.add_argument("--file-pattern",
                       default="*.json",
                       help="File pattern to match (default: *.json)")
    parser.add_argument("--create-docx",
                       action="store_true",
                       help="Create DOCX files alongside JSON files (requires python-docx)")
    parser.add_argument("--create-excel",
                       action="store_true",
                       help="Create Excel files alongside JSON files (requires pandas and openpyxl)")
    parser.add_argument("--create-all",
                       action="store_true",
                       help="Create all output formats (JSON, DOCX, and Excel)")
    parser.add_argument("--json-only",
                       action="store_true",
                       help="Create only JSON files (default behavior)")

    args = parser.parse_args()

    # Handle create-all flag
    if args.create_all:
        args.create_docx = True
        args.create_excel = True

    # Check for conflicting arguments
    if args.json_only and (args.create_docx or args.create_excel or args.create_all):
        print("Error: Cannot specify --json-only with other output format options")
        sys.exit(1)

    # Determine output formats
    create_docx = args.create_docx and not args.json_only
    create_excel = args.create_excel and not args.json_only

    # Check availability and warn if needed
    if create_docx and not DOCX_AVAILABLE:
        print("Warning: DOCX output requested but python-docx is not available.")
        print("Install with: pip install python-docx")
        print("Continuing without DOCX output...")
        create_docx = False

    if create_excel and not EXCEL_AVAILABLE:
        print("Warning: Excel output requested but pandas/openpyxl is not available.")
        print("Install with: pip install pandas openpyxl")
        print("Continuing without Excel output...")
        create_excel = False

    # Get the workspace root (assumes script is run from project root or scripts folder)
    if Path.cwd().name == "scripts":
        workspace_root = Path.cwd().parent
    else:
        workspace_root = Path.cwd()

    # Set up paths
    input_dir = workspace_root / args.input_dir
    output_dir = workspace_root / args.output_dir

    # Validate input directory
    if not input_dir.exists():
        print(f"Error: Input directory {input_dir} does not exist")
        sys.exit(1)

    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)

    # Find all JSON files in the input directory
    json_files = list(input_dir.glob(args.file_pattern))

    # Filter out hidden files and .keep files
    json_files = [f for f in json_files if not f.name.startswith('.') and f.name != '.keep']

    if not json_files:
        print(f"No JSON files found in {input_dir}")
        sys.exit(1)

    print(f"Found {len(json_files)} JSON files to process")
    print(f"Output directory: {output_dir}")
    print(f"Output structure:")
    print(f"  JSON files: {output_dir}/json/")
    if create_docx:
        print(f"  DOCX files: {output_dir}/docx/")
    if create_excel:
        print(f"  Excel files: {output_dir}/excel/")

    # Create output format summary
    output_formats = ["JSON"]
    if create_docx:
        output_formats.append("DOCX")
    if create_excel:
        output_formats.append("Excel")

    print(f"Output formats: {' + '.join(output_formats)}")
    if len(output_formats) > 1:
        print("Note: Each format will be created in separate folders")
    print("-" * 50)

    # Process each file
    total_records = 0
    for json_file in json_files:
        record_count = split_facility_file(json_file, output_dir, create_docx, create_excel)
        total_records += record_count

    # Create final summary
    output_formats = ["JSON"]
    if create_docx:
        output_formats.append("DOCX")
    if create_excel:
        output_formats.append("Excel")

    print("-" * 50)
    print(f"Processing complete!")
    print(f"Total files processed: {len(json_files)}")
    print(f"Total records created: {total_records}")
    print(f"Output formats: {' + '.join(output_formats)}")
    print(f"Output structure:")
    print(f"  JSON files location: {output_dir}/json/")
    if create_docx:
        print(f"  DOCX files location: {output_dir}/docx/")
    if create_excel:
        print(f"  Excel files location: {output_dir}/excel/")
    if not (create_docx or create_excel):
        print(f"  Base location: {output_dir}")

if __name__ == "__main__":
    main()
